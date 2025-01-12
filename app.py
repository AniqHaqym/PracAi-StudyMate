import streamlit as st
from jamaibase import JamAI, protocol as p
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from PyPDF2 import PdfReader

# Initialize JamAI client
jamai = JamAI(api_key="", project_id="")

# Initialize session state
if 'generated_content' not in st.session_state:
    st.session_state['generated_content'] = {
        'output': None,
        'study_topic': None,
        'has_generated': False
    }
    
# Add page number to session state
if 'page_number' not in st.session_state:
    st.session_state['page_number'] = 1

# Updated page navigation function
def update_page(new_page):
    """Update the current page number in session state"""
    st.session_state['page_number'] = new_page

def process_input():
    """Callback for processing input and updating session state"""
    if st.session_state.pdf_input and st.session_state.topic_input:
        try:
            # Extract text from PDF
            learning_material_text = extract_text_from_pdf(st.session_state.pdf_input)
            
            # Use JamAI API to generate study materials
            completion = jamai.add_table_rows(
                "action",
                p.RowAddRequest(
                    table_id="study-mate-final",
                    data=[{"learning_material": learning_material_text, "study_topic": st.session_state.topic_input}],
                    stream=False
                )
            )

            if completion.rows:
                output_row = completion.rows[0].columns
                # Update session state
                st.session_state['generated_content'] = {
                    'output': {
                        'study_plan': output_row.get("study_plan").text if output_row.get("study_plan") else 'N/A',
                        'summarized_notes': output_row.get("summarized_notes").text if output_row.get("summarized_notes") else 'N/A',
                        'quiz_questions': output_row.get("quiz_questions").text if output_row.get("quiz_questions") else 'N/A',
                        'supplementary_resources': output_row.get("supplementary_resources").text if output_row.get("supplementary_resources") else 'N/A'
                    },
                    'study_topic': st.session_state.topic_input,
                    'has_generated': True
                }
            else:
                st.error("⚠️ Failed to get a response. Please try again.")
        except Exception as e:
            st.error(f"❌ An error occurred: {e}")

def create_document():
    """Create and return a Word document with only the selected study materials"""
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Study Materials: {st.session_state["generated_content"]["study_topic"]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Only add sections that are selected
    if st.session_state.show_study_plan:
        doc.add_heading('Study Plan', 1)
        doc.add_paragraph(st.session_state['generated_content']['output']['study_plan'].strip())
        doc.add_page_break()
    
    if st.session_state.show_summarized_notes:
        doc.add_heading('Summarized Notes', 1)
        doc.add_paragraph(st.session_state['generated_content']['output']['summarized_notes'].strip())
        doc.add_page_break()
    
    if st.session_state.show_quiz_questions:
        doc.add_heading('Example Questions with Answers', 1)
        doc.add_paragraph(st.session_state['generated_content']['output']['quiz_questions'].strip())
        doc.add_page_break()
    
    if st.session_state.show_supplementary_resources:
        doc.add_heading('Supplementary Resources', 1)
        doc.add_paragraph(st.session_state['generated_content']['output']['supplementary_resources'].strip())
        doc.add_page_break()
    
    return doc

def get_document_bytes():
    """Get the document as bytes for download"""
    doc = create_document()
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf = PdfReader(pdf_file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

# Set up the Streamlit app
st.set_page_config(
    page_title="StudyMate",
    page_icon="📚",
    layout="centered"
)

# Updated CSS with theme-aware styling
st.markdown("""
    <style>
    /* Base theme-aware styles */
    :root {
        --text-color: rgb(49, 51, 63);
        --background-color: rgb(255, 255, 255);
    }

    /* Dark mode overrides */
    @media (prefers-color-scheme: dark) {
        :root {
            --text-color: rgb(250, 250, 250);
            --background-color: rgb(17, 17, 17);
        }
    }
    
    /* General container styling */
    .stApp {
        background: var(--background-color);
        color: var(--text-color);
    }
    
    /* Container styling */
    .css-1uhh5gq, .css-1dm3nw7, .css-1s12q1z {
        background-color: rgba(255, 255, 255, 0.1) !important;
        border-radius: 0.5rem;
        padding: 1.5rem;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.12);
        margin: 1rem 0;
        border: 1px solid rgba(128, 128, 128, 0.2);
    }
    
    /* Text styling */
    h1, h2, h3, label, .stMarkdown {
        color: var(--text-color) !important;
    }
    
    /* Button styling */
    .stButton > button {
        width: 100%;
        background-color: #3f51b5 !important;
        color: white !important;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
        border: none;
        font-weight: 600;
    }
    
    .stButton > button:hover {
        background-color: #1a237e !important;
        color: white !important;
    }
    
    /* File uploader styling */
    .stFileUploader {
        border: 2px dashed rgba(128, 128, 128, 0.4);
        border-radius: 10px;
        padding: 1rem;
        background-color: rgba(255, 255, 255, 0.05);
    }
    
    /* Text area styling */
    .stTextArea > div > div > textarea {
        background-color: var(--background-color) !important;
        color: var(--text-color) !important;
        border: 1px solid rgba(128, 128, 128, 0.4);
        border-radius: 0.5rem;
    }
    
    /* Checkbox styling */
    .stCheckbox > label {
        color: var(--text-color) !important;
    }
    
    /* Error and warning messages */
    .stAlert {
        background-color: rgba(255, 255, 255, 0.1) !important;
        color: var(--text-color) !important;
        border: 1px solid rgba(128, 128, 128, 0.2);
    }
    
    /* Navigation buttons */
    [data-testid="stHorizontalBlock"] button {
        background-color: #3f51b5 !important;
        color: white !important;
        min-width: 120px;
    }

    /* Download button styling */
    .stDownloadButton > button {
        background-color: #4CAF50 !important;
        color: white !important;
        min-width: 120px;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
        border: none;
        font-weight: 600;
    }
    
    .stDownloadButton > button:hover {
        background-color: #388E3C !important;
    }
    </style>
""", unsafe_allow_html=True)

# Set up the UI for input
st.title("📚 StudyMate - Your Personal AI Study Assistant")

with st.container():
    st.markdown("### 📄 Upload Learning Material and Study Topic")
    
    # File upload with custom styling
    st.file_uploader(
        "Upload Learning Material (PDF format)",
        type="pdf",
        key="pdf_input",
        help="Drag and drop your PDF file here"
    )

    # Topic input
    st.text_area(
        "✍️ Enter Topic/Keyword",
        key="topic_input",
        help="What would you like to study about?",
        placeholder="Enter the topic or keywords you want to study..."
    )

    # Options section
    st.markdown("### 🎯 Choose Study Options")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.checkbox("📋 Show Study Plan", value=True, key="show_study_plan")
        st.checkbox("📝 Show Summarized Notes", value=True, key="show_summarized_notes")
    
    with col2:
        st.checkbox("❓ Show Quiz Questions", value=True, key="show_quiz_questions")
        st.checkbox("📚 Show Resources", value=True, key="show_supplementary_resources")

    # Process button
    if st.button("🚀 Generate Study Materials", use_container_width=True, on_click=process_input):
        if not (st.session_state.pdf_input and st.session_state.topic_input):
            st.error("📢 Please upload a PDF and enter a topic/keyword.")
        if not any([st.session_state.show_study_plan, 
                   st.session_state.show_summarized_notes,
                   st.session_state.show_quiz_questions,
                   st.session_state.show_supplementary_resources]):
            st.warning("⚠️ Please select at least one study option.")

# Page Navigation Logic
if st.session_state['generated_content']['has_generated']:
    # Get the selected options and create a pages list
    pages = []
    if st.session_state.show_study_plan:
        pages.append("Study Plan")
    if st.session_state.show_summarized_notes:
        pages.append("Summarized Notes")
    if st.session_state.show_quiz_questions:
        pages.append("Quiz Questions")
    if st.session_state.show_supplementary_resources:
        pages.append("Supplementary Resources")
    
    total_pages = len(pages)
    
    if total_pages > 0:
        # Ensure the page_number doesn't go out of range
        current_page = min(max(st.session_state['page_number'] - 1, 0), total_pages - 1)
        
        # Show content for the current page
        page_name = pages[current_page]
        if page_name == "Study Plan":
            st.markdown("## 📋 Study Plan")
            st.markdown(st.session_state['generated_content']['output']['study_plan'])
        elif page_name == "Summarized Notes":
            st.markdown("## 📝 Summarized Notes")
            st.markdown(st.session_state['generated_content']['output']['summarized_notes'])
        elif page_name == "Quiz Questions":
            st.markdown("## ❓ Quiz Questions with Answers")
            st.markdown(st.session_state['generated_content']['output']['quiz_questions'])
        elif page_name == "Supplementary Resources":
            st.markdown("## 📚 Supplementary Resources")
            st.markdown(st.session_state['generated_content']['output']['supplementary_resources'])

        # Navigation Buttons
        col1, col2 = st.columns(2)
        with col1:
            if current_page > 0:
                if st.button("← Previous", key="prev_button"):
                    update_page(current_page)
        with col2:
            if current_page < total_pages - 1:
                if st.button("Next →", key="next_button"):
                    update_page(current_page + 2)
                
    # Download button
    sanitized_topic = "".join(c for c in st.session_state['generated_content']['study_topic'].lower() 
                             if c.isalnum() or c in (' ', '_')).replace(" ", "_")
    
    st.download_button(
        label="📥 Download Study Materials as Word Document",
        data=get_document_bytes(),
        file_name=f"study_materials_{sanitized_topic}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
